#!/usr/bin/env python
# encoding=utf-8

#导入依赖的所有库
import os
import shutil
import re
import time
import datetime
import numpy
import pandas
import docx
import xlrd
import xlwt
import openpyxl
import selenium
from selenium import webdriver


##定义一系列函数

# Catalogue Test:检测目录是否被别人做了
def catatest(catali,tsl=0,uname=None,pw=None,hide=True,pd=False):
        '''
        catalogue test
        pass a 'catalogue list' parameter,return a 'kezuo and yiyou'2-dimension list
        catali must be a list, consisting of the topics that you want to search on the site.
        '''
        import time
        import datetime
        import numpy as np
        import selenium
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys

        option = webdriver.ChromeOptions()
        option.add_argument("headless")
        
        t1=datetime.datetime.now()
        
        if hide==True:
            sch = webdriver.Chrome(options=option)
        else:
            sch = webdriver.Chrome()
        #sch=webdriver.Chrome()

        site=r'https://workbench.qyresearch.com/login'
        sch.get(site)
        
        time.sleep(tsl*np.random.rand()+1.002)
        
        sch.find_element_by_id('username').clear()
        sch.find_element_by_id('username').send_keys(uname)
        sch.find_element_by_id('password').clear()
        sch.find_element_by_id('password').send_keys(pw)
        sch.find_element_by_id('password').send_keys(Keys.ENTER)
        
        time.sleep(tsl*np.random.rand()+1.002)

        kezuo=[]
        yiyou=[]

        n=0
        for i in catali:
            inp=sch.find_element_by_xpath(r'//*[@id="keywords"]')
            inp.send_keys(i)
            
            time.sleep(tsl*np.random.rand()+1.002)
            
            sebu=sch.find_element_by_xpath(r'/html/body/app-root/app-layout/div/section/div/app-global-search/div/div[2]/form/div[3]/button')
            sebu.click()
            
            time.sleep(tsl*np.random.rand()+1.002)
            
            try: 
                a=sch.find_element_by_xpath(r'/html/body/app-root/app-layout/div/section/div/app-global-search/div/app-common-table/div/div[2]')
                if a.text=='暂无数据 No Data':
                    #print(a.text)
                    kezuo.append(i)
                    n+=1
                else:
                    pass
            except:
                b1=sch.find_element_by_xpath(r'/html/body/app-root/app-layout/div/section/div/app-global-search/div/app-common-table/*')
                b2=b1.text
                yiyou.append(b2)
            inp.clear()
        sch.quit()
        #print('可做的哦：','一共',n,'个','\n','----','\n',kezuo)
        
        if pd==True:
            import pandas as pd
            shuru=pd.Series(catali,name='输入')
            kezuo=pd.Series(kezuo,name='可做')
            resu=pd.concat([shuru,kezuo],axis=1,join='outer')
        else:
            resu=[kezuo,yiyou]
        
        t2=datetime.datetime.now()
        print('总共用时： ',t2-t1)
        
        return resu

# Find files with Regular Expression：正则表达式搜索本地文件
def findfile(item,fdir,match=False):
    '''
    Find the file you need in a directory.
    Regular Expression is supported.
    '''
    import os
    import re
    rs=[]
    for i,j,k in os.walk(fdir):
        for na in k:
            if match==True:
                regitem=re.compile(item)
                d=re.match(regitem,na)
                if d is not None:
                    rs.append(os.path.join(i,na))
            else:
                regitem=re.compile(item)
                d=re.search(regitem,na)
                if d is not None:
                    rs.append(os.path.join(i,na))
    if rs==[]:
        print(r"Can't find anything!")
    return rs


#搜索文件，并复制所有搜索结果到指定文件夹
def findcpall(item,fdir,destination,match=False):
    import os
    import re
    import shutil
    #rs=[]
    if match==True:
        for i,j,k in os.walk(fdir):
            for na in k:
                regitem=re.compile(item)
                d=re.match(regitem,na)
                if d is not None:
                    rsdir=os.path.join(i,na)
                    shutil.copy(rsdir,destination)
    else:
        for i,j,k in os.walk(fdir):
            for na in k:
                regitem=re.compile(item)
                d=re.search(regitem,na)
                if d is not None:
                    rsdir=os.path.join(i,na)
                    shutil.copy(rsdir,destination)
    print('copy finished')

#接到Sample预处理，在server上搜索节选库和产品库，找到相关文件全部复制：
def samprepare(item,badir=r'D:\HZ.SK\MissionAccomplished\Sample节选',mtch=False):
    '''
    item is a str indicating the name of the study objective.
    '''
    import os
    import re
    import shutil
    import datetime

    t1=datetime.datetime.now()
    #print('start time: ',t1)


    #创建item专属的目录
    itemdir=str(badir)+os.sep+str(item)
    os.mkdir(itemdir)
    
    #创建 Final目录
    final=itemdir+os.sep+str(r'Final')
    os.mkdir(final)
    
    #在item专属目录下，创建多个参考目录
    
    #reprefdir=itemdir+os.sep+str(r'ReportRef')
    #os.mkdir(reprefdir)
        
    samrefdir=itemdir+os.sep+str(r'SampleRef')
    os.mkdir(samrefdir)
        
    #contrefdir=itemdir+os.sep+str(r'ContentRef')
    #os.mkdir(contrefdir)
    
    prodrefdir=itemdir+os.sep+str(r'ProductRef')
    os.mkdir(prodrefdir)
    
    fdir=[r'\\server\1Report',
         r'\\server\2Sample节选库',
         r'\\server\3Content目录库',
         r'\\server\4Product产品库']
    
    def findcpall(item,fdir,destination,match=False):
        import os
        import re
        import shutil
        #rs=[]
        if match==True:
            for i,j,k in os.walk(fdir):
                for na in k:
                    regitem=re.compile(item)
                    d=re.match(regitem,na)
                    if d is not None:
                        rsdir=os.path.join(i,na)
                        shutil.copy(rsdir,destination)
        else:
            for i,j,k in os.walk(fdir):
                for na in k:
                    regitem=re.compile(item)
                    d=re.search(regitem,na)
                    if d is not None:
                        rsdir=os.path.join(i,na)
                        shutil.copy(rsdir,destination)
        #print('copy finished')
    
    #findcpall(item,fdir[0],reprefdir,match=mtch)
    findcpall(item,fdir[1],samrefdir,match=mtch)
    #findcpall(item,fdir[2],contrefdir,match=mtch)
    findcpall(item,fdir[3],prodrefdir,match=mtch)
    
    t2=datetime.datetime.now()
    #print('finish time: ',t2)
    print('总共用时： ',t2-t1)
#
#接到Project预处理，在server上搜索已有报告和sample，找到相关文件全部复制：
def projpp(item,datainfo=True,mtch=False):
    '''
    item is a str indicating the name of the study objective,and do not use regular expression.
    datainfo=True, means you're searching for data information, else you're getting chapter information.
    '''
    import os
    import re
    import shutil
    import datetime

    t1=datetime.datetime.now()
    #print('start time: ',t1)

    if datainfo==True:
        badir=r'D:\HZ.SK\MissionAccomplished\Project 数据'
    else:
        badir=r'D:\HZ.SK\MissionAccomplished\Project 章节'
    
    #创建item专属的目录
    itemdir=str(badir)+os.sep+str(item)
    os.mkdir(itemdir)
    
    #创建 Final目录
    final=itemdir+os.sep+str(r'Final')
    os.mkdir(final)
    
    #在item专属目录下，创建多个参考目录
    
    reprefdir=itemdir+os.sep+str(r'ReportRef')
    os.mkdir(reprefdir)
        
    samrefdir=itemdir+os.sep+str(r'SampleRef')
    os.mkdir(samrefdir)
        
    #contrefdir=itemdir+os.sep+str(r'ContentRef')
    #os.mkdir(contrefdir)
    
    #prodrefdir=itemdir+os.sep+str(r'ProductRef')
    #os.mkdir(prodrefdir)
    
    fdir=[r'\\server\1Report',
         r'\\server\2Sample节选库',
         r'\\server\3Content目录库',
         r'\\server\4Product产品库']
    
    def findcpall(item,fdir,destination,match=False):
        import os
        import re
        import shutil
        #rs=[]
        if match==True:
            for i,j,k in os.walk(fdir):
                for na in k:
                    regitem=re.compile(item)
                    d=re.match(regitem,na)
                    if d is not None:
                        rsdir=os.path.join(i,na)
                        shutil.copy(rsdir,destination)
        else:
            for i,j,k in os.walk(fdir):
                for na in k:
                    regitem=re.compile(item)
                    d=re.search(regitem,na)
                    if d is not None:
                        rsdir=os.path.join(i,na)
                        shutil.copy(rsdir,destination)
        #print('copy finished')
    
    findcpall(item,fdir[0],reprefdir,match=mtch)
    findcpall(item,fdir[1],samrefdir,match=mtch)
    #findcpall(item,fdir[2],contrefdir,match=mtch)
    #findcpall(item,fdir[3],prodrefdir,match=mtch)
    
    t2=datetime.datetime.now()
    #print('finish time: ',t2)
    print('总共用时： ',t2-t1)
#

#打开若干网站
def opsite(siteli,tsl=0):
    '''
    open a list of websites.
    siteli must be a list, each element of which is one of the websites that you want to open in one single web browser.
    '''
    import time
    import numpy as np
    from selenium import webdriver
    osite=webdriver.Chrome()
    n=1
    for i in siteli:
        osite.get(i)
        time.sleep(tsl*np.random.rand()+0.002)
        osite.execute_script("window.open()")
        handles=osite.window_handles
        osite.switch_to.window(handles[n])
        n+=1
    osite.close()
    osite.switch_to.window(handles[0])
    
#简单Google搜索：按照对象列表com+关键词列表kw组合搜索。
def simplegoo(com,kw=[''],tsl=0):
    '''
    Google companies plus the keyword given.
    com and kw must be lists, excluding NaN elements.
    tsl is the time sleep level when opening each tab.
    '''
    import time
    import numpy as np
    from selenium import webdriver
    from selenium.webdriver.common.keys import Keys
    
    wh=r"https://www.google.com/"
    js=r"window.open('https://www.google.com/')"
    
    sch=webdriver.Chrome()
    
    li=[]
    for i in com:
        for j in kw:
            b=str(i)+r' '+str(j)
            li.append(b)
    n=1
    for i in li:
            a=i
            sch.get(wh)
            sch.find_element_by_name("q").send_keys(a)
            sch.find_element_by_name("q").send_keys(Keys.ENTER)
            time.sleep(tsl*np.random.rand()+0.002)
            sch.execute_script(js)
            handles=sch.window_handles
            sch.switch_to.window(handles[n])
            n+=1
    sch.close()
    sch.switch_to.window(handles[0])


#自动填入ProductName和CompanyList
#做目录的Excel里边不同的产品行数相差56。
#产品中文名称：B2；Product Name：C2（c58） #company 1 的位置：C4(c60)
#这个fillcali适合从Excel读取一个DataFrame，而不是直接输入DataFrame。
def fillcali(cali,indir,savedir,start_prod='C2',start_com='C4',cellstep=56):
    '''
    Fill the category list into the Excel.
    catali must be a pandas.DataFrame where the first product starts at the second column, 
    and of course, the first column would be the "id" column.
    '''
    import openpyxl as ox
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter, column_index_from_string
    
    wb=ox.load_workbook(indir)
    ws=wb['产品信息']
    
    nali=cali.columns
    nali=list(nali)
    nali.pop(0)
    
    def onefill(prodbcell,combcell,n):
        '''
        n的范围从0开始，表示列表nali当中元素的索引。
        n starts from 0, indicating to the index of the list nali.
        '''
        prodbcell.value=nali[n]
        
        comli=cali[nali[n]]
        comli=comli.iloc[range(12)].dropna()
        comli=list(comli)
        le=len(comli)
        for i in range(le):
            combcell.offset(i,0).value=comli[i]
    
    prodbcell=ws[start_prod]
    combcell=ws[start_com]
    
    for i in range(len(nali)):
        onefill(prodbcell,combcell,i)
        prodbcell=prodbcell.offset(cellstep,0)
        combcell=combcell.offset(cellstep,0)
    
    wb.save(savedir)
#
#这个fillcalia适合直接输入一个pandas.DataFrame，而不是从Excel读取一个DataFrame：
def fillcalia(cali,indir,savedir,start_prod='C2',start_com='C4',cellstep=56):
    '''
    Fill the category list into the Excel.
    "catali" must be a pandas.DataFrame where the first product starts at the first column, 
    and of course, the first column would not be the "id" column.
    "indir" could be xlsm, yet the "savedir" must be xlsx, and, certainly, VBA will not be removed.
    '''
    import openpyxl as ox
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter, column_index_from_string
    
    wb=ox.load_workbook(indir,keep_vba=True)
    ws=wb['产品信息']
    
    nali=cali.columns
    nali=list(nali)
    if nali[0] == 0:
        nali.pop(0)
    #print(cali.columns)
    #print(nali)
    
    def onefill(prodbcell,combcell,n):
        '''
        n starts from 0, indicating to the index of the list nali.
        '''
        prodbcell.value=nali[n]
        
        comli=cali[nali[n]]
        comli=comli.iloc[range(12)].dropna()
        comli=list(comli)
        le=len(comli)
        for i in range(le):
            combcell.offset(i,0).value=comli[i]
    
    prodbcell=ws[start_prod]
    combcell=ws[start_com]
    
    for i in range(len(nali)):
        onefill(prodbcell,combcell,i)
        prodbcell=prodbcell.offset(cellstep,0)
        combcell=combcell.offset(cellstep,0)
    
    wb.save(savedir)    
#
  
#函数-获取docx文档概览：
def getscan(indir,depth=1,accu=True):
    '''
    Get scan of a Microsoft Word file.
    indir is the inputing directory,and depth indicates how deep you want to know.
    e.g.with depth=3, you'll get all the Heading 1, Heading 2 as well as Heading 3 of the file.
    and,if accu is False, you'll get Heading 3 only.
    '''
    resu=[]
    import docx
    f=docx.Document(indir)
    for i in f.paragraphs:
        if accu==False:
            name_left='Heading '
            name_right=str(depth)
            if i.style.name==name_left+name_right:
                resu.append(i.text)
        else:
            j=1
            for j in range(1,depth+1):
                name_left='Heading '
                name_right=str(j)
                if i.style.name==name_left+name_right:
                    resu.append(i.text)
                    j+=1
    return resu


#对比两个Word文档(docx)结构的不同，返回一个pandas.DataFrame:
def dcompare(sdir1,sdir2,dep=1,accum=True):
    def getscan(indir,depth=1,accu=True):
        resu=[]
        import docx
        f=docx.Document(indir)
        for i in f.paragraphs:
            if accu==False:
                name_left='Heading '
                name_right=str(depth)
                if i.style.name==name_left+name_right:
                    resu.append(i.text)
            else:
                j=1
                for j in range(1,depth+1):
                    name_left='Heading '
                    name_right=str(j)
                    if i.style.name==name_left+name_right:
                        resu.append(i.text)
                        j+=1
        return resu
    
    s1=getscan(sdir1,depth=dep,accu=accum)
    s2=getscan(sdir2,depth=dep,accu=accum)
    import pandas as pd
    s1=pd.Series(s1)
    s2=pd.Series(s2)
    dresu=pd.concat([s1,s2],axis=1,join='outer')
    return dresu
    

##定制化的继承类 webdriver.Chrome
from selenium import webdriver
class Scha(webdriver.Chrome):
    #简单搜索
    def simplegoo(self,com,kw=[''],tsl=0):
        '''
        Google companies plus the keyword given.
        com and kw must be lists, excluding NaN elements.
        tsl is the time sleep level when opening each tab.
        If you want to add extensions when starting a class Sch, use the following code :
        {
        from selenium import webdriver
        option = webdriver.ChromeOptions()
        #if you want to hide the web broswer window,put "headless" into:
        #option.add_argument("headless") 
        #if you want to add extensions, put the directory of the "crx" file into:
        option.add_extension(r'D:\HZ.SK\tools\Chrome插件\ClusterTabManager\2.2.3_0.crx')
        from skoffice import Sch
        }
        '''
        import time
        import numpy as np
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys

        wh=r"https://www.google.com/"
        js=r"window.open('https://www.google.com/')"

        li=[]
        for i in com:
            for j in kw:
                b=str(i)+r' '+str(j)
                li.append(b)
        n=1
        for i in li:
                a=i
                self.get(wh)
                self.find_element_by_name("q").send_keys(a)
                self.find_element_by_name("q").send_keys(Keys.ENTER)
                time.sleep(tsl*np.random.rand()+0.002)
                self.execute_script(js)
                handles=self.window_handles
                self.switch_to.window(handles[n])
                n+=1
        self.close()
        self.switch_to.window(handles[0])
    
    def simpledu(self,com,kw=[''],tsl=0):
        '''
        Google废了的时候用百度，比必应还厉害呢！
        '''
        import time
        import numpy as np
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys
        
        wh=r"https://www.baidu.com/"
        js=r"window.open('https://www.baidu.com/')"
        
        li=[]
        for i in com:
            for j in kw:
                b=str(i)+r' '+str(j)
                li.append(b)
        n=1
        for i in li:
                a=i
                self.get(wh)
                self.find_element_by_id("kw").send_keys(a)
                self.find_element_by_id("kw").send_keys(Keys.ENTER)
                
                time.sleep(tsl*np.random.rand()+0.502)
                
                self.execute_script(js)
                handles=self.window_handles
                self.switch_to.window(handles[n])
                n+=1
        self.close()
        self.switch_to.window(handles[0])
    
    def simplebi(self,com,kw=[''],tsl=0):
        '''
        其实必应搜索并没有百度厉害
        '''
        import time
        import numpy as np
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys
        
        wh=r"https://bing.com/"
        js=r"window.open('https://bing.com/')"
        
        li=[]
        for i in com:
            for j in kw:
                b=str(i)+r' '+str(j)
                li.append(b)
        n=1
        for i in li:
                a=i
                self.get(wh)
                self.find_element_by_id("sb_form_q").send_keys(a)
                self.find_element_by_id("sb_form_q").send_keys(Keys.ENTER)
                time.sleep(tsl*np.random.rand()+0.45)
                self.find_element_by_id("est_en").click()
                time.sleep(tsl*np.random.rand()+0.2)
                self.execute_script(js)
                handles=self.window_handles
                self.switch_to.window(handles[n])
                n+=1
        self.close()
        self.switch_to.window(handles[0])
    
    
    #打开若干网站
    def opsite(self,siteli,tsl=0):
        '''
        open a list of websites.
        siteli must be a list, each element of which is one of the websites that you want to open in one single web browser.
        '''
        import time
        import numpy as np
        from selenium import webdriver
        
        n=1
        for i in siteli:
            self.get(i)
            time.sleep(tsl*np.random.rand()+0.002)
            self.execute_script("window.open()")
            handles=self.window_handles
            self.switch_to.window(handles[n])
            n+=1
        self.close()
        self.switch_to.window(handles[0])
    #
    
    #在capterra上搜索得到vendor list：
    def getvdcap(self,capsiteli,tsl=0):
        '''
        capsiteli is the list of product links from capterra.
        '''
        import time
        import re
        import numpy as np
        import pandas as pd
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys
        #wh=r'https://www.capterra.com/'
        #js=r'window.open("https://www.capterra.com/")'
        
        def getvenli(self):
            venli=[]
            cvendors=self.find_elements_by_css_selector(r'div.cell div.cell h3.epsilon')
            for i in cvendors:
                i=i.text
                reg=re.search(r'\b[^(by\s)].+',i).span()
                i=i[reg[0]:reg[1]]
                venli.append(i)
            return venli
        
        resu=pd.DataFrame([])
        
        for i in capsiteli:
            self.get(i)
            time.sleep(tsl*np.random.rand()+0.2)
            self.find_element_by_xpath(r'//*[@id="sort_options_select"]').click()
            self.find_element_by_xpath(r'//*[@id="sort_options_select"]/option[3]').click()
            na=self.find_element_by_xpath(r'/html/body/div[1]/div[1]/div/div/div/h1').text
            vdlist=getvenli(self)
            vdlist=pd.Series(vdlist,name=na)
            resu=pd.concat([resu,vdlist],axis=1,join='outer')
        self.close()
        return resu
    #



##自动生成新闻稿攻略