#!/usr/bin/env python
# encoding='utf-8'

# 将名为"vardict-whatdata.xlsx"的Excel文件放在脚本同一目录下。输入参数，可得目标文档的结构。

topic_fake=input(r'=>Step1 type in the project name:')
indir_fake=input(r'=>Step2 type in the Word Docx directory to read:')
outdir_fake1=input(r'=>Step3 type in the outputing directory:')


# vardict_dir=r'D:\HZ.SK\MissionAccomplished\Project 数据\vardict-whatdata.xlsx'
vardict_dir=r'./vardict-whatdata.xlsx'

#读取Word中每个6级标题,提取关键信息,输出信息矩阵.
import re
import os
#import docx
#import openpyxl as ox
#import numpy as np
import pandas as pd
#import skaida as ska
#from skoffice import findfile

out_name=r'SimpleInfo-'+topic_fake+r'.xlsx'
outdir_fake=os.path.join(outdir_fake1,out_name)
print(r'outputing directory:',outdir_fake)

vardict=pd.read_excel(vardict_dir,sheet_name=r'vardict')
#vardict=pd.DataFrame(columns=['subj','share','by','year','reg'])

# for i in vardict:
    # for j in vardict[i].dropna():
        # f1=re.compile(j)
        # print(f1)

##

#函数-结构化解析列表中的元素：
def readstrli(instrli,vardict=pd.read_excel(vardict_dir,sheet_name=r'vardict')):
    ##
    def readstr(instr):
        def oneread(tli,tstr):
            bli=[]
            for i in tli:
                acom=re.compile(i)
                a=re.search(acom,tstr)
                if a != None:
                    b=a.group(0)
                    bli.append(b)
                else:
                    bli.append('nothing')
            for j in bli:
                if j != 'nothing':
                    b_final=j
                    return b_final
        vardict_col=list(vardict.columns)
        subinfo = []
        for i in vardict_col:
            tli1=vardict.loc[:,i].dropna()
            tli1=list(tli1)
            c=oneread(tli1,instr)
            subinfo.append(c)
        return subinfo
        print(subinfo)
    #
    infomat=[]
    for i in instrli:
        i_fake=str(i)
        subinfo1=readstr(i_fake)
        infomat.append(subinfo1)
    infomat=pd.DataFrame(infomat,columns=vardict.columns)
    return infomat
#
#函数-获取docx文档概览：
#def getscan(indir,depth=1,accu=True):
#    '''
#    Get scan of a Microsoft Word file.
#    indir is the inputing directory,and depth indicates how deep you want to know.
#    e.g.with depth=3, you'll get all the Heading 1, Heading 2 as well as Heading 3 of the file.
#    and,if accu is False, you'll get Heading 3 only.
#    '''
#    resu=[]
#    import docx
#    f=docx.Document(indir)
#    for i in f.paragraphs:
#        if accu==False:
#            name_left='Heading '
#            name_right=str(depth)
#            if i.style.name==name_left+name_right:
#                resu.append(i.text)
#        else:
#            j=1
#            for j in range(1,depth+1):
#                name_left='Heading '
#                name_right=str(j)
#                if i.style.name==name_left+name_right:
#                    resu.append(i.text)
#                    j+=1
#    return resu
#
def drawtf(indir):
    from docx import Document
    indoc=Document(indir)
    tfli=[]
    for i in indoc.paragraphs:
        if i.style.name=='Table Style':
            tfli.append(r'Table '+i.text)
        elif i.style.name=='Figure Style':
            tfli.append(r'Figure '+i.text)
        else:
            pass
    return tfli

#函数-正则表达式匹配列表中的元素：
#def regtestli(item,li,match=False):
#    '''
#    Test elements in a list with Regular Expression.
#    Extract elements that work well with a Regular Expression (item) from a list(li), and return a result list.
#    '''
#    import re
#    c=re.compile(item)
#    n=0
#    rs=[]
#    for i in li:
#        if match == True:
#            b=re.match(c,i)
#            if b is not None:
#                rs.append(li[n])
#        else:
#            b=re.search(c,i)
#            if b is not None:
#                rs.append(li[n])
#        n+=1
#    return rs
#

#print(os.path.abspath(os.curdir))

#t1=getscan(indir_fake,depth=6,accu=False)
t1=drawtf(indir_fake)

#t2=regtestli(r'\d{4}(\s*(to|-|-)\s*\d{4})?',t1,match=False)
#print(r'图表数目：',len(t2))

print(r'图表数目：',len(t1))

t3=readstrli(t1)
t3['original']=t1
#print('t3 head-3 SCAN:',t3.head(3),r"t3's shape:",t3.shape)

t3.to_excel(outdir_fake,sheet_name=r'Catalogue')

###
